#!/usr/bin/env python3
"""
Construye el módulo de vigencia cartográfica de IPT.

Cruza:
- reportes mensuales IPT;
- histórico anual;
- catálogo de versiones cartográficas;
- actos manuales;
- comparaciones espaciales opcionales en GeoJSON.

El resultado distingue entre:
- Actualizado;
- Probablemente actualizado;
- Revisión necesaria;
- Desactualizado;
- Sin cartografía.

Una alerta temporal no prueba por sí sola que el shape esté desactualizado.
La clasificación "Desactualizado" exige evidencia explícita, ausencia de una
zona esperada o una comparación espacial bajo el umbral configurado.
"""

from __future__ import annotations

import csv
import io
import json
import re
import sys
import unicodedata
from collections import defaultdict
from datetime import date, datetime
from pathlib import Path
from typing import Any
from urllib.parse import urlparse
from zoneinfo import ZoneInfo

import requests
from docx import Document
from docx.shared import Cm, Pt
from shapely.geometry import shape
from shapely.ops import unary_union


ROOT = Path(__file__).resolve().parents[1]
CONFIG_FILE = ROOT / "config" / "cartografia_ipt.json"
IPT_FILE = ROOT / "data" / "ipt_reportes.js"
HISTORIC_FILE = ROOT / "data" / "historicos.js"
OUTPUT_FILE = ROOT / "data" / "vigencia_cartografica.js"
DOCS_DIR = ROOT / "documentos" / "vigencia"
CSV_DIR = ROOT / "consolidados" / "vigencia"
TIMEZONE = ZoneInfo("America/Santiago")
REQUEST_TIMEOUT = 45

CARTOGRAPHIC_CONFIRMED = (
    "plano", "zonificacion", "zonificación", "limite urbano", "límite urbano",
    "uso de suelo", "vialidad estructurante", "sector modificado",
    "reemplaza el plano", "incorpora zona", "crea zona", "amplia el limite",
    "amplía el límite", "seccional"
)
CARTOGRAPHIC_PROBABLE = (
    "modificacion", "modificación", "enmienda", "rectificacion",
    "rectificación", "actualizacion", "actualización"
)
NON_FINAL_STATES = ("en tramitacion", "en tramitación", "consulta", "inicio")


class VigenciaError(RuntimeError):
    pass


def clean(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def normalized(value: Any) -> str:
    text = unicodedata.normalize("NFKD", clean(value))
    text = "".join(char for char in text if not unicodedata.combining(char))
    text = text.lower()
    return re.sub(r"[^a-z0-9]+", " ", text).strip()


def slug(*values: Any) -> str:
    text = "-".join(normalized(value).replace(" ", "-") for value in values if clean(value))
    return re.sub(r"-+", "-", text).strip("-") or "instrumento"


def parse_date(value: Any) -> date | None:
    text = clean(value)
    if not text:
        return None

    match = re.search(r"(\d{4})[-/](\d{1,2})[-/](\d{1,2})", text)
    if match:
        try:
            return date(int(match.group(1)), int(match.group(2)), int(match.group(3)))
        except ValueError:
            return None

    match = re.search(r"\b(19|20)\d{2}\b", text)
    if match:
        try:
            return date(int(match.group(0)), 1, 1)
        except ValueError:
            return None

    return None


def iso_or_empty(value: Any) -> str:
    parsed = parse_date(value)
    return parsed.isoformat() if parsed else ""


def load_js(path: Path, prefix: str) -> Any:
    raw = path.read_text(encoding="utf-8").strip()
    if not raw.startswith(prefix) or not raw.endswith(";"):
        raise VigenciaError(f"Formato inválido en {path.name}.")
    return json.loads(raw[len(prefix):-1])


def load_catalog() -> dict[str, Any]:
    if not CONFIG_FILE.exists():
        raise VigenciaError("No existe config/cartografia_ipt.json.")
    value = json.loads(CONFIG_FILE.read_text(encoding="utf-8"))
    if not isinstance(value, dict):
        raise VigenciaError("El catálogo cartográfico debe ser un objeto JSON.")
    value.setdefault("instrumentos", [])
    value.setdefault("actos_manuales", [])
    value.setdefault("parametros", {})
    return value


def infer_cartographic_impact(act: dict[str, Any]) -> str:
    explicit = act.get("requiere_cambio_cartografico")
    if explicit is True:
        return "confirmado"
    if explicit is False:
        return "no"

    text = normalized(" ".join([
        clean(act.get("tipo_acto")),
        clean(act.get("tipo_ipt")),
        clean(act.get("titulo")),
        clean(act.get("resumen")),
    ]))

    if any(normalized(keyword) in text for keyword in CARTOGRAPHIC_CONFIRMED):
        return "confirmado"
    if any(normalized(keyword) in text for keyword in CARTOGRAPHIC_PROBABLE):
        return "probable"
    return "no_determinado"


def act_identifier(act: dict[str, Any]) -> str:
    return clean(
        act.get("id")
        or act.get("numero")
        or act.get("titulo")
        or f"{act.get('tipo_acto', '')}-{act.get('fecha_publicacion', '')}"
    )


def extract_acts(
    ipt_reports: list[dict[str, Any]],
    historical_reports: list[dict[str, Any]],
    manual_acts: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    acts: list[dict[str, Any]] = []

    for report in ipt_reports:
        for change in report.get("cambios", []) or []:
            act = {
                "id": slug(
                    "ipt",
                    report.get("periodo"),
                    change.get("region"),
                    change.get("comuna"),
                    change.get("numero"),
                    change.get("acto"),
                ),
                "region": clean(change.get("region")),
                "comuna": clean(change.get("comuna")),
                "tipo_ipt": clean(change.get("tipo_ipt")) or "IPT",
                "tipo_acto": clean(change.get("acto")) or "Acto IPT",
                "numero": clean(change.get("numero")),
                "fecha_publicacion": iso_or_empty(change.get("fecha_publicacion")),
                "estado": clean(change.get("estado")) or "Otro",
                "titulo": clean(change.get("acto")) or clean(change.get("tipo_ipt")),
                "resumen": clean(change.get("resumen")),
                "fuente": clean(change.get("fuente")),
                "vigencia": clean(change.get("vigencia")),
                "zonas_afectadas": change.get("zonas_afectadas", []) or [],
                "archivo_geojson": clean(change.get("archivo_geojson")),
                "campo_zona": clean(change.get("campo_zona")),
                "zona_esperada": clean(change.get("zona_esperada")),
                "operacion": clean(change.get("operacion")),
                "origen": "Reporte IPT mensual",
            }
            act["impacto_cartografico"] = infer_cartographic_impact(act)
            acts.append(act)

    for report in historical_reports:
        for item in report.get("items", []) or []:
            if clean(item.get("modulo")) != "IPT":
                continue
            act = {
                "id": slug(
                    "historico",
                    item.get("periodo"),
                    item.get("region"),
                    item.get("comuna"),
                    item.get("numero"),
                    item.get("titulo"),
                ),
                "region": clean(item.get("region")),
                "comuna": clean(item.get("comuna")),
                "tipo_ipt": clean(item.get("tipo_norma")) or "IPT",
                "tipo_acto": clean(item.get("categoria")) or "Acto IPT",
                "numero": clean(item.get("numero")),
                "fecha_publicacion": iso_or_empty(item.get("fecha")),
                "estado": clean(item.get("estado")) or "Otro",
                "titulo": clean(item.get("titulo")),
                "resumen": clean(item.get("resumen")),
                "fuente": clean(item.get("fuente")),
                "vigencia": clean(item.get("implicancia")),
                "zonas_afectadas": item.get("zonas_afectadas", []) or [],
                "archivo_geojson": clean(item.get("archivo_geojson")),
                "campo_zona": clean(item.get("campo_zona")),
                "zona_esperada": clean(item.get("zona_esperada")),
                "operacion": clean(item.get("operacion")),
                "origen": "Histórico anual",
            }
            act["impacto_cartografico"] = infer_cartographic_impact(act)
            acts.append(act)

    for manual in manual_acts:
        act = {
            "id": clean(manual.get("id")) or slug(
                "manual",
                manual.get("region"),
                manual.get("comuna"),
                manual.get("numero"),
                manual.get("titulo"),
            ),
            "region": clean(manual.get("region")),
            "comuna": clean(manual.get("comuna")),
            "tipo_ipt": clean(manual.get("tipo_ipt")) or "IPT",
            "tipo_acto": clean(manual.get("tipo_acto")) or "Acto IPT",
            "numero": clean(manual.get("numero")),
            "fecha_publicacion": iso_or_empty(manual.get("fecha_publicacion")),
            "estado": clean(manual.get("estado")) or "Otro",
            "titulo": clean(manual.get("titulo")),
            "resumen": clean(manual.get("resumen")),
            "fuente": clean(manual.get("fuente")),
            "vigencia": clean(manual.get("vigencia")),
            "requiere_cambio_cartografico": manual.get("requiere_cambio_cartografico"),
            "zonas_afectadas": manual.get("zonas_afectadas", []) or [],
            "archivo_geojson": clean(manual.get("archivo_geojson")),
            "campo_zona": clean(manual.get("campo_zona")),
            "zona_esperada": clean(manual.get("zona_esperada")),
            "operacion": clean(manual.get("operacion")),
            "origen": "Catálogo manual",
        }
        act["impacto_cartografico"] = infer_cartographic_impact(act)
        acts.append(act)

    unique: dict[tuple[str, ...], dict[str, Any]] = {}
    for act in acts:
        key = (
            normalized(act.get("region")),
            normalized(act.get("comuna")),
            normalized(act.get("tipo_ipt")),
            normalized(act.get("numero")),
            normalized(act.get("fecha_publicacion")),
            normalized(act.get("titulo")),
        )
        current = unique.get(key)
        if current is None or current.get("origen") != "Catálogo manual":
            unique[key] = act

    return sorted(
        unique.values(),
        key=lambda item: (
            item.get("fecha_publicacion", ""),
            item.get("titulo", ""),
        ),
    )


def read_json_reference(reference: str) -> dict[str, Any] | None:
    ref = clean(reference)
    if not ref:
        return None

    parsed = urlparse(ref)
    if parsed.scheme in {"http", "https"}:
        response = requests.get(ref, timeout=REQUEST_TIMEOUT)
        response.raise_for_status()
        return response.json()

    path = Path(ref)
    if not path.is_absolute():
        path = ROOT / path
    if not path.exists():
        return None
    return json.loads(path.read_text(encoding="utf-8"))


def geojson_union(
    data: dict[str, Any],
    field: str = "",
    expected_value: str = "",
):
    features = data.get("features", []) if isinstance(data, dict) else []
    geometries = []

    for feature in features:
        if not isinstance(feature, dict) or not feature.get("geometry"):
            continue

        if field and expected_value:
            properties = feature.get("properties", {}) or {}
            if normalized(properties.get(field)) != normalized(expected_value):
                continue

        try:
            geometries.append(shape(feature["geometry"]))
        except Exception:
            continue

    if not geometries:
        return None
    return unary_union(geometries)


def spatial_check(
    instrument: dict[str, Any],
    act: dict[str, Any],
    parameters: dict[str, Any],
) -> dict[str, Any] | None:
    base_ref = clean(instrument.get("archivo_geojson"))
    act_ref = clean(act.get("archivo_geojson"))
    if not base_ref or not act_ref:
        return None

    base_data = read_json_reference(base_ref)
    act_data = read_json_reference(act_ref)
    if not base_data or not act_data:
        return {
            "estado": "no_ejecutada",
            "coincidencia_porcentaje": None,
            "observacion": "No fue posible abrir uno de los archivos GeoJSON."
        }

    field = clean(act.get("campo_zona")) or clean(instrument.get("campo_zona"))
    expected = clean(act.get("zona_esperada"))

    modification_geometry = geojson_union(act_data)
    if modification_geometry is None or modification_geometry.is_empty:
        return {
            "estado": "no_ejecutada",
            "coincidencia_porcentaje": None,
            "observacion": "La geometría de la modificación está vacía."
        }

    target_geometry = geojson_union(base_data, field, expected) if expected else geojson_union(base_data)
    if target_geometry is None or target_geometry.is_empty:
        return {
            "estado": "no_incorporado",
            "coincidencia_porcentaje": 0.0,
            "observacion": (
                f"No se encontró la zona esperada {expected!r} en el shape."
                if expected
                else "El shape no contiene geometrías comparables."
            )
        }

    denominator = modification_geometry.area
    if denominator <= 0:
        return {
            "estado": "no_ejecutada",
            "coincidencia_porcentaje": None,
            "observacion": "La geometría de la modificación no tiene superficie."
        }

    coverage = modification_geometry.intersection(target_geometry).area / denominator
    upper = float(parameters.get("umbral_actualizado", 0.95))
    lower = float(parameters.get("umbral_desactualizado", 0.70))

    if coverage >= upper:
        status = "incorporado"
    elif coverage < lower:
        status = "no_incorporado"
    else:
        status = "revision"

    return {
        "estado": status,
        "coincidencia_porcentaje": round(coverage * 100, 2),
        "observacion": (
            f"Cobertura de la geometría modificada por la zona esperada: "
            f"{coverage * 100:.2f}%."
        )
    }


def token_matches(act: dict[str, Any], values: list[Any]) -> bool:
    act_tokens = {
        normalized(act.get("id")),
        normalized(act.get("numero")),
        normalized(act.get("titulo")),
        normalized(act_identifier(act)),
    }
    act_tokens.discard("")

    for value in values or []:
        candidate = normalized(value)
        if not candidate:
            continue
        if candidate in act_tokens:
            return True
        if any(candidate in token or token in candidate for token in act_tokens):
            return True
    return False


def relevant_to_instrument(act: dict[str, Any], instrument: dict[str, Any]) -> bool:
    if normalized(act.get("comuna")) != normalized(instrument.get("comuna")):
        return False

    instrument_type = normalized(instrument.get("tipo_ipt"))
    act_type = normalized(act.get("tipo_ipt"))
    if not instrument_type or not act_type:
        return True
    return instrument_type in act_type or act_type in instrument_type or "prc" in {instrument_type, act_type}


def classify_instrument(
    instrument: dict[str, Any],
    acts: list[dict[str, Any]],
    parameters: dict[str, Any],
) -> dict[str, Any]:
    version_date = (
        parse_date(instrument.get("fecha_version_cartografica"))
        or parse_date(instrument.get("fecha_instrumento_base"))
    )
    incorporated = instrument.get("actos_incorporados", []) or []
    explicit_missing = instrument.get("actos_no_incorporados", []) or []
    consider_pending = bool(parameters.get("considerar_actos_en_tramitacion", False))

    alerts: list[dict[str, Any]] = []
    timeline: list[dict[str, Any]] = []
    spatial_results: list[dict[str, Any]] = []
    pending_acts: list[dict[str, Any]] = []
    confirmed_missing = False

    base_date = iso_or_empty(instrument.get("fecha_instrumento_base"))
    if base_date:
        timeline.append({
            "fecha": base_date,
            "tipo": "Instrumento base",
            "titulo": clean(instrument.get("nombre")) or clean(instrument.get("tipo_ipt")),
            "estado": "Vigente",
            "incorporacion": "base",
            "fuente": clean(instrument.get("fuente_instrumento")),
            "resumen": "Entrada en vigencia del instrumento base."
        })

    shape_date = iso_or_empty(instrument.get("fecha_version_cartografica"))
    if shape_date:
        timeline.append({
            "fecha": shape_date,
            "tipo": "Versión cartográfica",
            "titulo": "Shape o capa cartográfica analizada",
            "estado": "Disponible",
            "incorporacion": "shape",
            "fuente": clean(instrument.get("fuente_cartografia")),
            "resumen": clean(instrument.get("notas"))
        })

    for act in acts:
        act_date = parse_date(act.get("fecha_publicacion"))
        non_final = any(
            state in normalized(act.get("estado"))
            for state in NON_FINAL_STATES
        )
        impact = act.get("impacto_cartografico", "no_determinado")
        is_incorporated = token_matches(act, incorporated)
        is_explicit_missing = token_matches(act, explicit_missing)

        spatial = spatial_check(instrument, act, parameters)
        if spatial:
            spatial_results.append({
                "acto_id": act.get("id"),
                "acto": act_identifier(act),
                **spatial
            })
            if spatial.get("estado") == "incorporado":
                is_incorporated = True
            elif spatial.get("estado") == "no_incorporado":
                is_explicit_missing = True

        later_than_shape = (
            version_date is None
            or act_date is None
            or act_date > version_date
        )
        potentially_relevant = impact in {"confirmado", "probable"}

        if is_explicit_missing and potentially_relevant:
            confirmed_missing = True
            alerts.append({
                "nivel": "alto",
                "tipo": "Acto no incorporado",
                "acto_id": act.get("id"),
                "mensaje": (
                    f"{act_identifier(act)} tiene evidencia de no estar "
                    "incorporado en la versión cartográfica."
                )
            })
        elif (
            potentially_relevant
            and later_than_shape
            and not is_incorporated
            and (consider_pending or not non_final)
        ):
            pending_acts.append(act)
            alerts.append({
                "nivel": "alto" if impact == "confirmado" else "medio",
                "tipo": "Acto posterior al shape",
                "acto_id": act.get("id"),
                "mensaje": (
                    f"{act_identifier(act)} es posterior a la versión "
                    "cartográfica y no tiene evidencia de incorporación."
                )
            })

        if is_explicit_missing:
            incorporation = "no_incorporado"
        elif is_incorporated:
            incorporation = "incorporado"
        elif not later_than_shape:
            incorporation = "probablemente_incorporado"
        else:
            incorporation = "sin_verificar"

        timeline.append({
            "fecha": clean(act.get("fecha_publicacion")),
            "tipo": clean(act.get("tipo_acto")),
            "titulo": clean(act.get("titulo")) or act_identifier(act),
            "numero": clean(act.get("numero")),
            "estado": clean(act.get("estado")),
            "incorporacion": incorporation,
            "impacto_cartografico": impact,
            "fuente": clean(act.get("fuente")),
            "resumen": clean(act.get("resumen")),
            "zonas_afectadas": act.get("zonas_afectadas", []) or [],
            "archivo_geojson": clean(act.get("archivo_geojson")),
            "zona_esperada": clean(act.get("zona_esperada")),
        })

    has_cartography = bool(
        clean(instrument.get("archivo_geojson"))
        or clean(instrument.get("fuente_cartografia"))
        or clean(instrument.get("fecha_version_cartografica"))
    )
    validated_spatial = [
        item for item in spatial_results if item.get("estado") == "incorporado"
    ]

    if not has_cartography:
        status = "Sin cartografía"
        confidence = "alta"
        summary = "Existe una línea normativa, pero no hay una versión cartográfica registrada para comparar."
    elif confirmed_missing:
        status = "Desactualizado"
        confidence = "alta"
        summary = "Existe evidencia explícita o espacial de actos no incorporados."
    elif pending_acts:
        status = "Revisión necesaria"
        confidence = "alta" if any(
            act.get("impacto_cartografico") == "confirmado" for act in pending_acts
        ) else "media"
        summary = "Hay actos posteriores con posible efecto cartográfico sin evidencia suficiente de incorporación."
    elif validated_spatial:
        status = "Actualizado"
        confidence = "alta"
        summary = "Las comparaciones espaciales disponibles cumplen el umbral configurado."
    else:
        status = "Probablemente actualizado"
        confidence = "media"
        summary = "No se detectan actos cartográficos pendientes, pero falta validación espacial completa."

    timeline.sort(key=lambda item: item.get("fecha", ""))

    map_overlays = [
        {
            "id": act.get("id"),
            "titulo": act_identifier(act),
            "archivo_geojson": clean(act.get("archivo_geojson")),
            "incorporacion": next(
                (
                    event.get("incorporacion")
                    for event in timeline
                    if event.get("titulo") == (clean(act.get("titulo")) or act_identifier(act))
                ),
                "sin_verificar",
            ),
            "zona_esperada": clean(act.get("zona_esperada")),
        }
        for act in acts
        if clean(act.get("archivo_geojson"))
    ]

    return {
        "id": clean(instrument.get("id")) or slug(
            instrument.get("region"),
            instrument.get("comuna"),
            instrument.get("tipo_ipt"),
        ),
        "region": clean(instrument.get("region")),
        "comuna": clean(instrument.get("comuna")),
        "tipo_ipt": clean(instrument.get("tipo_ipt")) or "IPT",
        "nombre": clean(instrument.get("nombre")) or clean(instrument.get("tipo_ipt")),
        "fecha_instrumento_base": base_date,
        "fecha_version_cartografica": shape_date,
        "fuente_cartografia": clean(instrument.get("fuente_cartografia")),
        "archivo_geojson": clean(instrument.get("archivo_geojson")),
        "campo_zona": clean(instrument.get("campo_zona")),
        "zonas_presentes": instrument.get("zonas_presentes", []) or [],
        "estado_alerta": status,
        "confianza": confidence,
        "resumen_alerta": summary,
        "actos_posteriores_pendientes": len(pending_acts),
        "alertas": alerts,
        "linea_tiempo": timeline,
        "comparaciones_espaciales": spatial_results,
        "mapa": {
            "base_geojson": clean(instrument.get("archivo_geojson")),
            "capas_modificaciones": map_overlays,
        },
        "notas": clean(instrument.get("notas")),
    }


def virtual_instruments(
    acts: list[dict[str, Any]],
    configured: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    configured_keys = {
        (normalized(item.get("comuna")), normalized(item.get("tipo_ipt")))
        for item in configured
    }
    groups: dict[tuple[str, str, str], list[dict[str, Any]]] = defaultdict(list)

    for act in acts:
        commune = clean(act.get("comuna"))
        if not commune:
            continue
        key = (
            normalized(commune),
            normalized(act.get("tipo_ipt")),
        )
        if key in configured_keys:
            continue
        groups[(
            clean(act.get("region")),
            commune,
            clean(act.get("tipo_ipt")) or "IPT",
        )].append(act)

    result = []
    for (region, commune, instrument_type), group_acts in groups.items():
        instrument = {
            "id": slug("sin-cartografia", region, commune, instrument_type),
            "region": region,
            "comuna": commune,
            "tipo_ipt": instrument_type,
            "nombre": instrument_type,
            "fecha_instrumento_base": "",
            "fecha_version_cartografica": "",
            "fuente_cartografia": "",
            "archivo_geojson": "",
            "campo_zona": "",
            "zonas_presentes": [],
            "actos_incorporados": [],
            "actos_no_incorporados": [],
            "notas": "Instrumento creado automáticamente desde actos IPT; falta registrar el shape.",
        }
        result.append((instrument, group_acts))

    return result


def create_csv(instruments: list[dict[str, Any]]) -> Path:
    CSV_DIR.mkdir(parents=True, exist_ok=True)
    output = CSV_DIR / "Alertas_vigencia_cartografica.csv"

    headers = [
        "Región", "Comuna", "Tipo IPT", "Instrumento",
        "Fecha instrumento base", "Fecha versión cartográfica",
        "Estado", "Confianza", "Actos pendientes", "Resumen",
        "Fuente cartografía",
    ]

    with output.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle, delimiter=";")
        writer.writerow(headers)
        for item in instruments:
            writer.writerow([
                item.get("region", ""),
                item.get("comuna", ""),
                item.get("tipo_ipt", ""),
                item.get("nombre", ""),
                item.get("fecha_instrumento_base", ""),
                item.get("fecha_version_cartografica", ""),
                item.get("estado_alerta", ""),
                item.get("confianza", ""),
                item.get("actos_posteriores_pendientes", 0),
                item.get("resumen_alerta", ""),
                item.get("fuente_cartografia", ""),
            ])

    return output


def create_word(instruments: list[dict[str, Any]], generated_at: str) -> Path:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    output = DOCS_DIR / "Reporte_vigencia_cartografica_IPT.docx"

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)

    document.styles["Normal"].font.name = "Outfit"
    document.styles["Normal"].font.size = Pt(10)
    document.styles["Title"].font.name = "Outfit"
    document.styles["Title"].font.size = Pt(22)
    document.styles["Heading 1"].font.name = "Outfit"
    document.styles["Heading 1"].font.size = Pt(15)
    document.styles["Heading 2"].font.name = "Outfit"
    document.styles["Heading 2"].font.size = Pt(12)

    document.add_heading("Reporte de vigencia cartográfica de IPT", 0)
    document.add_paragraph(
        f"Departamento de Estudios Inmobiliarios · Transsa\nGenerado: {generated_at}"
    ).runs[0].bold = True

    counts = defaultdict(int)
    for instrument in instruments:
        counts[instrument.get("estado_alerta", "Sin clasificación")] += 1

    document.add_heading("Síntesis ejecutiva", level=1)
    document.add_paragraph(
        f"Se evaluaron {len(instruments)} instrumentos o líneas normativas. "
        f"Desactualizados: {counts['Desactualizado']}; "
        f"revisión necesaria: {counts['Revisión necesaria']}; "
        f"sin cartografía: {counts['Sin cartografía']}."
    )

    for instrument in instruments:
        document.add_heading(
            " · ".join(filter(None, [
                instrument.get("comuna"),
                instrument.get("tipo_ipt"),
                instrument.get("nombre"),
            ])),
            level=1,
        )

        for label, key in [
            ("Región", "region"),
            ("Estado", "estado_alerta"),
            ("Confianza", "confianza"),
            ("Fecha del instrumento base", "fecha_instrumento_base"),
            ("Fecha de la versión cartográfica", "fecha_version_cartografica"),
            ("Fuente cartográfica", "fuente_cartografia"),
        ]:
            value = instrument.get(key)
            if value:
                paragraph = document.add_paragraph()
                paragraph.add_run(f"{label}: ").bold = True
                paragraph.add_run(str(value))

        document.add_paragraph(instrument.get("resumen_alerta", ""))

        timeline = instrument.get("linea_tiempo", []) or []
        if timeline:
            document.add_heading("Línea de tiempo", level=2)
            for event in timeline:
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.add_run(
                    f"{event.get('fecha') or 'Sin fecha'} · "
                    f"{event.get('tipo') or 'Acto'} · "
                    f"{event.get('titulo') or ''}"
                ).bold = True
                incorporation = clean(event.get("incorporacion"))
                if incorporation:
                    paragraph.add_run(f" — {incorporation.replace('_', ' ')}")
                if event.get("resumen"):
                    document.add_paragraph(event["resumen"])

    note = document.add_paragraph(
        "La evaluación temporal es una alerta de control. La conclusión jurídica "
        "y cartográfica definitiva requiere revisar los actos oficiales y, cuando "
        "corresponda, comparar espacialmente la modificación con el shape."
    )
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(8.5)

    document.save(output)
    return output


def main() -> int:
    catalog = load_catalog()
    ipt_reports = load_js(IPT_FILE, "window.IPT_REPORTES = ")
    historical = load_js(HISTORIC_FILE, "window.HISTORICOS = ")

    if not isinstance(ipt_reports, list) or not isinstance(historical, list):
        raise VigenciaError("Las bases IPT e histórica deben ser listas.")

    parameters = catalog.get("parametros", {}) or {}
    acts = extract_acts(
        ipt_reports,
        historical,
        catalog.get("actos_manuales", []) or [],
    )

    configured = catalog.get("instrumentos", []) or []
    instruments: list[dict[str, Any]] = []

    for instrument in configured:
        matching = [
            act for act in acts
            if relevant_to_instrument(act, instrument)
        ]
        instruments.append(
            classify_instrument(instrument, matching, parameters)
        )

    for instrument, matching in virtual_instruments(acts, configured):
        instruments.append(
            classify_instrument(instrument, matching, parameters)
        )

    instruments.sort(
        key=lambda item: (
            normalized(item.get("region")),
            normalized(item.get("comuna")),
            normalized(item.get("tipo_ipt")),
        )
    )

    generated_at = datetime.now(TIMEZONE).strftime("%Y-%m-%d %H:%M")
    csv_file = create_csv(instruments)
    word_file = create_word(instruments, generated_at)

    counts = defaultdict(int)
    for instrument in instruments:
        counts[instrument.get("estado_alerta", "")] += 1

    output = {
        "fecha_generacion": generated_at,
        "resumen": {
            "instrumentos": len(instruments),
            "actualizados": counts["Actualizado"],
            "probablemente_actualizados": counts["Probablemente actualizado"],
            "revision_necesaria": counts["Revisión necesaria"],
            "desactualizados": counts["Desactualizado"],
            "sin_cartografia": counts["Sin cartografía"],
        },
        "instrumentos": instruments,
        "word_url": word_file.relative_to(ROOT).as_posix(),
        "csv_url": csv_file.relative_to(ROOT).as_posix(),
        "nota_metodologica": (
            "La evaluación temporal no reemplaza la validación espacial ni "
            "jurídica del instrumento. Solo se clasifica como desactualizado "
            "cuando existe evidencia explícita o espacial."
        ),
    }

    OUTPUT_FILE.write_text(
        "window.VIGENCIA_CARTOGRAFICA = "
        + json.dumps(output, ensure_ascii=False, indent=2)
        + ";\n",
        encoding="utf-8",
    )

    print(f"Instrumentos evaluados: {len(instruments)}")
    print(f"Revisión necesaria: {counts['Revisión necesaria']}")
    print(f"Desactualizados: {counts['Desactualizado']}")
    print(f"Sin cartografía: {counts['Sin cartografía']}")
    print(f"Salida: {OUTPUT_FILE.relative_to(ROOT).as_posix()}")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except (VigenciaError, requests.RequestException, json.JSONDecodeError) as exc:
        print(f"Error: {exc}", file=sys.stderr)
        raise SystemExit(2)
