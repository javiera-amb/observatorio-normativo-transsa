#!/usr/bin/env python3
"""
Actualiza mensualmente el módulo de Instrumentos de Planificación Territorial.

Cobertura:
- período calendario anterior;
- 16 regiones de Chile;
- fuentes web oficiales;
- PRC, PRI, PRM, planes seccionales, enmiendas, límites urbanos,
  postergaciones, PROT y otros actos que cambien o tramiten normativa
  territorial.

Salida:
- data/ipt_reportes.js
- Word mensual
- CSV consolidado compatible con Excel

La revisión automática amplía la cobertura, pero no reemplaza la validación
profesional de los actos oficiales ni garantiza descubrir publicaciones que
no estén indexadas públicamente.
"""

from __future__ import annotations

import csv
import json
import os
import re
import sys
from datetime import date, datetime
from pathlib import Path
from typing import Any
from zoneinfo import ZoneInfo

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from openai import OpenAI


ROOT = Path(__file__).resolve().parents[1]
DATA_FILE = ROOT / "data" / "ipt_reportes.js"
DOCUMENTS_DIR = ROOT / "documentos" / "ipt"
CONSOLIDATED_DIR = ROOT / "consolidados" / "ipt"
TIMEZONE = ZoneInfo("America/Santiago")
MODEL = os.environ.get("OPENAI_MODEL", "gpt-5-mini")

REGIONS = [
    "Arica y Parinacota",
    "Tarapacá",
    "Antofagasta",
    "Atacama",
    "Coquimbo",
    "Valparaíso",
    "Metropolitana de Santiago",
    "O'Higgins",
    "Maule",
    "Ñuble",
    "Biobío",
    "La Araucanía",
    "Los Ríos",
    "Los Lagos",
    "Aysén",
    "Magallanes y de la Antártica Chilena",
]

ALLOWED_DOMAINS = [
    "diariooficial.interior.gob.cl",
    "minvu.gob.cl",
    "ide.minvu.cl",
    "gob.cl",
    "mma.gob.cl",
    "eae.mma.gob.cl",
    "sea.gob.cl",
    "bcn.cl",
]

SCHEMA_FIELDS = [
    "region",
    "comuna",
    "territorio",
    "tipo_ipt",
    "acto",
    "numero",
    "fecha_publicacion",
    "estado",
    "resumen",
    "vigencia",
    "fuente",
]


class AutomationError(RuntimeError):
    pass


def previous_month(reference: date) -> tuple[int, int]:
    if reference.month == 1:
        return reference.year - 1, 12
    return reference.year, reference.month - 1


def month_bounds(year: int, month: int) -> tuple[date, date]:
    start = date(year, month, 1)
    if month == 12:
        next_month = date(year + 1, 1, 1)
    else:
        next_month = date(year, month + 1, 1)
    end = date.fromordinal(next_month.toordinal() - 1)
    return start, end


def month_name(year: int, month: int) -> str:
    names = [
        "enero", "febrero", "marzo", "abril", "mayo", "junio",
        "julio", "agosto", "septiembre", "octubre", "noviembre",
        "diciembre",
    ]
    return f"{names[month - 1]} de {year}"


def clean(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def parse_json_object(raw: str) -> dict[str, Any]:
    text = raw.strip()
    text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s*```$", "", text)

    try:
        result = json.loads(text)
    except json.JSONDecodeError:
        start = text.find("{")
        end = text.rfind("}")
        if start < 0 or end <= start:
            raise AutomationError("La respuesta no contiene JSON válido.")
        result = json.loads(text[start:end + 1])

    if not isinstance(result, dict):
        raise AutomationError("La respuesta debe ser un objeto JSON.")
    return result


def client() -> OpenAI:
    if not os.environ.get("OPENAI_API_KEY"):
        raise AutomationError("Falta OPENAI_API_KEY.")
    return OpenAI()


def research_region(
    api: OpenAI,
    region: str,
    start: date,
    end: date,
) -> list[dict[str, Any]]:
    prompt = f"""
Eres analista territorial del Departamento de Estudios Inmobiliarios de
Transsa. Investiga en la web actualizaciones oficiales ocurridas entre
{start.isoformat()} y {end.isoformat()} para la región de {region}, Chile.

Busca actos publicados, aprobados, promulgados, iniciados o sometidos a
consulta que afecten instrumentos de planificación territorial:

- planes reguladores comunales, intercomunales y metropolitanos;
- planes seccionales y enmiendas;
- límites urbanos;
- postergaciones de permisos;
- planes regionales de ordenamiento territorial;
- modificaciones, rectificaciones, actualizaciones y aprobaciones;
- procedimientos de evaluación ambiental estratégica de estos instrumentos.

Revisa todas las comunas de la región en la medida en que existan fuentes
oficiales indexadas. Usa únicamente fuentes oficiales o institucionales.
No incluyas simples noticias sin un acto o trámite identificable.

Devuelve exclusivamente JSON válido:
{{
  "cambios": [
    {{
      "region": "{region}",
      "comuna": "Comuna o vacío si es intercomunal/regional",
      "territorio": "Territorio afectado",
      "tipo_ipt": "PRC, PRI, PRM, seccional, límite urbano, PROT, etc.",
      "acto": "Tipo de acto o etapa",
      "numero": "Número del acto si existe",
      "fecha_publicacion": "YYYY-MM-DD o vacío",
      "estado": "Vigente | En tramitación | Consulta | Otro",
      "resumen": "Resumen concreto del cambio",
      "vigencia": "Qué está vigente y qué aún no",
      "fuente": "URL completa de la fuente oficial"
    }}
  ],
  "advertencias_cobertura": [
    "Fuentes o comunas que no pudieron verificarse completamente"
  ]
}}

No inventes URLs ni actos. Si no encuentras cambios, devuelve cambios vacío.
""".strip()

    response = api.responses.create(
        model=MODEL,
        input=prompt,
        tools=[
            {
                "type": "web_search",
                "filters": {"allowed_domains": ALLOWED_DOMAINS},
                "search_context_size": "high",
                "user_location": {
                    "type": "approximate",
                    "country": "CL",
                    "region": region,
                    "timezone": "America/Santiago",
                },
            }
        ],
        store=False,
    )

    result = parse_json_object(response.output_text)
    changes = result.get("cambios", [])
    if not isinstance(changes, list):
        return []

    cleaned: list[dict[str, Any]] = []
    for item in changes:
        if not isinstance(item, dict):
            continue

        record = {field: clean(item.get(field, "")) for field in SCHEMA_FIELDS}
        record["region"] = record["region"] or region

        if record["estado"] not in {
            "Vigente", "En tramitación", "Consulta", "Otro"
        }:
            record["estado"] = "Otro"

        if not record["resumen"] or not record["fuente"]:
            continue

        cleaned.append(record)

    return cleaned


def deduplicate(records: list[dict[str, Any]]) -> list[dict[str, Any]]:
    output: list[dict[str, Any]] = []
    seen: set[tuple[str, ...]] = set()

    for record in records:
        key = (
            clean(record.get("region")).lower(),
            clean(record.get("comuna")).lower(),
            clean(record.get("tipo_ipt")).lower(),
            clean(record.get("acto")).lower(),
            clean(record.get("numero")).lower(),
            clean(record.get("fecha_publicacion")).lower(),
        )
        if key in seen:
            continue
        seen.add(key)
        output.append(record)

    return sorted(
        output,
        key=lambda item: (
            item.get("region", ""),
            item.get("comuna", ""),
            item.get("fecha_publicacion", ""),
        ),
    )


def load_reports() -> list[dict[str, Any]]:
    raw = DATA_FILE.read_text(encoding="utf-8").strip()
    prefix = "window.IPT_REPORTES = "

    if not raw.startswith(prefix) or not raw.endswith(";"):
        raise AutomationError("Formato inválido en data/ipt_reportes.js.")

    reports = json.loads(raw[len(prefix):-1])
    if not isinstance(reports, list):
        raise AutomationError("La base IPT no es una lista.")
    return reports


def save_reports(reports: list[dict[str, Any]]) -> None:
    reports.sort(key=lambda item: item.get("periodo", ""), reverse=True)
    DATA_FILE.write_text(
        "window.IPT_REPORTES = "
        + json.dumps(reports, ensure_ascii=False, indent=2)
        + ";\n",
        encoding="utf-8",
    )


def create_csv(period: str, records: list[dict[str, Any]]) -> Path:
    CONSOLIDATED_DIR.mkdir(parents=True, exist_ok=True)
    output = CONSOLIDATED_DIR / f"Consolidado_IPT_{period}.csv"

    headers = [
        "Región",
        "Comuna",
        "Territorio",
        "Tipo de IPT",
        "Acto o norma",
        "Número",
        "Fecha de publicación",
        "Estado",
        "Resumen",
        "Vigencia",
        "Fuente oficial",
    ]

    with output.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle, delimiter=";")
        writer.writerow(headers)

        for record in records:
            writer.writerow([
                record.get("region", ""),
                record.get("comuna", ""),
                record.get("territorio", ""),
                record.get("tipo_ipt", ""),
                record.get("acto", ""),
                record.get("numero", ""),
                record.get("fecha_publicacion", ""),
                record.get("estado", ""),
                record.get("resumen", ""),
                record.get("vigencia", ""),
                record.get("fuente", ""),
            ])

    return output


def create_word(
    period: str,
    period_label: str,
    records: list[dict[str, Any]],
    generated_at: str,
) -> Path:
    DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)
    output = DOCUMENTS_DIR / f"Reporte_actualizaciones_IPT_{period}.docx"

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)

    document.styles["Normal"].font.name = "Outfit"
    document.styles["Normal"].font.size = Pt(10)
    document.styles["Title"].font.name = "Outfit"
    document.styles["Title"].font.size = Pt(22)
    document.styles["Heading 1"].font.name = "Outfit"
    document.styles["Heading 1"].font.size = Pt(15)
    document.styles["Heading 2"].font.name = "Outfit"
    document.styles["Heading 2"].font.size = Pt(12)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("Reporte mensual de actualizaciones IPT")

    subtitle = document.add_paragraph()
    subtitle.add_run(
        f"Departamento de Estudios Inmobiliarios · Transsa\n"
        f"Período revisado: {period_label}\n"
        f"Generado: {generated_at}"
    ).bold = True

    document.add_heading("Síntesis ejecutiva", level=1)

    communes = len({r.get("comuna") for r in records if r.get("comuna")})
    regions = len({r.get("region") for r in records if r.get("region")})

    if records:
        document.add_paragraph(
            f"Se identificaron {len(records)} cambios o actuaciones relevantes "
            f"en {communes} comunas y {regions} regiones."
        )
    else:
        document.add_paragraph(
            "No se identificaron cambios de IPT respaldados por fuentes "
            "oficiales indexadas durante el período."
        )

    current_region = None
    for record in records:
        if record["region"] != current_region:
            current_region = record["region"]
            document.add_heading(current_region, level=1)

        heading = " · ".join(
            value for value in [
                record.get("comuna") or record.get("territorio"),
                record.get("tipo_ipt"),
                record.get("acto"),
            ] if value
        )
        document.add_heading(heading or "Actualización territorial", level=2)

        for label, key in [
            ("Número", "numero"),
            ("Fecha", "fecha_publicacion"),
            ("Estado", "estado"),
            ("Vigencia", "vigencia"),
        ]:
            value = record.get(key, "")
            if value:
                paragraph = document.add_paragraph()
                paragraph.add_run(f"{label}: ").bold = True
                paragraph.add_run(value)

        paragraph = document.add_paragraph()
        paragraph.add_run("Resumen: ").bold = True
        paragraph.add_run(record.get("resumen", ""))

        paragraph = document.add_paragraph()
        paragraph.add_run("Fuente oficial: ").bold = True
        paragraph.add_run(record.get("fuente", ""))

    document.add_paragraph()
    note = document.add_paragraph(
        "La revisión automática depende de la indexación pública de las fuentes. "
        "Cada antecedente debe verificarse en el acto oficial antes de aplicarlo "
        "a un análisis normativo o inmobiliario."
    )
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(8.5)

    document.save(output)
    return output


def main() -> int:
    now = datetime.now(TIMEZONE)
    year, month = previous_month(now.date())
    start, end = month_bounds(year, month)
    period = f"{year:04d}-{month:02d}"
    label = month_name(year, month)

    reports = load_reports()
    if any(report.get("periodo") == period for report in reports):
        print(f"El período {period} ya está procesado.")
        return 0

    api = client()
    all_changes: list[dict[str, Any]] = []

    for region in REGIONS:
        print(f"Revisando {region}...")
        try:
            all_changes.extend(research_region(api, region, start, end))
        except Exception as exc:
            print(
                f"Advertencia en {region}: {type(exc).__name__}: {exc}",
                file=sys.stderr,
            )

    changes = deduplicate(all_changes)
    generated_at = now.strftime("%Y-%m-%d %H:%M")

    word_path = create_word(period, label, changes, generated_at)
    csv_path = create_csv(period, changes)

    communes = len({r.get("comuna") for r in changes if r.get("comuna")})
    regions = len({r.get("region") for r in changes if r.get("region")})

    if changes:
        summary = (
            f"Se identificaron {len(changes)} cambios o actuaciones relevantes "
            f"en {communes} comunas y {regions} regiones durante {label}."
        )
    else:
        summary = (
            f"No se identificaron cambios de IPT respaldados por fuentes "
            f"oficiales indexadas durante {label}."
        )

    report = {
        "periodo": period,
        "titulo": f"Actualizaciones IPT · {label.capitalize()}",
        "fecha_generacion": generated_at,
        "resumen_ejecutivo": summary,
        "cambios": changes,
        "word_url": word_path.relative_to(ROOT).as_posix(),
        "csv_url": csv_path.relative_to(ROOT).as_posix(),
        "excel_url": "",
        "alcance": "Revisión nacional por 16 regiones",
        "nota_cobertura": (
            "La revisión automática depende de la indexación pública de las "
            "fuentes oficiales y requiere validación profesional."
        ),
    }

    reports.append(report)
    save_reports(reports)

    print(f"Período procesado: {period}")
    print(f"Cambios: {len(changes)}")
    print(f"Word: {report['word_url']}")
    print(f"CSV: {report['csv_url']}")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except AutomationError as exc:
        print(f"Error: {exc}", file=sys.stderr)
        raise SystemExit(2)
