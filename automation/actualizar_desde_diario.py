#!/usr/bin/env python3
"""
Actualiza el Observatorio Normativo Urbano a partir de la edición más reciente
del Diario Oficial de la República de Chile.

Flujo:
1. Descarga el sumario oficial.
2. Detecta la edición y sus publicaciones PDF.
3. Usa OpenAI para seleccionar las materias pertinentes.
4. Descarga y extrae el texto de los PDF oficiales.
5. Genera registros estructurados para el portal.
6. Crea un Word diario descargable.
7. Actualiza data/reportes.js sin duplicar registros.

Variables de entorno:
- OPENAI_API_KEY: obligatoria cuando existe una nueva edición por analizar.
- OPENAI_MODEL: opcional. Valor por defecto: gpt-5-mini.
"""

from __future__ import annotations

import io
import json
import os
import re
import sys
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path
from typing import Any
from urllib.parse import urljoin
from zoneinfo import ZoneInfo

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from openai import OpenAI
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
DATA_FILE = ROOT / "data" / "reportes.js"
DOCUMENTS_DIR = ROOT / "documentos"

INDEX_URL = (
    "https://www.diariooficial.interior.gob.cl/"
    "edicionelectronica/index.php/bom.php"
)
TIMEZONE = ZoneInfo("America/Santiago")
MODEL = os.environ.get("OPENAI_MODEL", "gpt-5-mini")
REQUEST_TIMEOUT = 45
MAX_PDF_CHARS = 45_000
USER_AGENT = (
    "Mozilla/5.0 (compatible; ObservatorioNormativoTranssa/1.0; "
    "+https://www.diariooficial.interior.gob.cl/)"
)

MONTHS = {
    "enero": 1,
    "febrero": 2,
    "marzo": 3,
    "abril": 4,
    "mayo": 5,
    "junio": 6,
    "julio": 7,
    "agosto": 8,
    "septiembre": 9,
    "setiembre": 9,
    "octubre": 10,
    "noviembre": 11,
    "diciembre": 12,
}

RELEVANCE_CRITERIA = """
Selecciona publicaciones que tengan relación directa o material con alguna
de estas materias:

- planificación urbana, territorial o metropolitana;
- instrumentos de planificación territorial, PRC, PRI, PRM, planes
  seccionales, enmiendas, límites urbanos o declaratorias;
- usos de suelo, subdivisiones, loteos, urbanizaciones y aportes al espacio
  público;
- permisos, anteproyectos, recepciones, regularizaciones y Direcciones de
  Obras Municipales;
- LGUC, OGUC, normas técnicas, construcción, arquitectura, edificación,
  accesibilidad, seguridad, instalaciones o eficiencia energética;
- vivienda, subsidios y programas habitacionales cuando cambien reglas,
  montos, llamados o condiciones aplicables a proyectos;
- patrimonio arquitectónico, zonas típicas, monumentos, inmuebles de
  conservación o arqueología vinculada a obras;
- evaluación ambiental de proyectos urbanos, inmobiliarios, industriales o
  de infraestructura con efectos territoriales, viales, patrimoniales o de
  construcción;
- movilidad, vialidad, expropiaciones, infraestructura urbana y espacio
  público;
- normas regionales o comunales que puedan afectar análisis inmobiliarios.

Ante duda razonable, incluye la publicación. Excluye materias claramente
ajenas, como pesca, tipos de cambio, nombramientos internos o procedimientos
administrativos sin efecto territorial.
""".strip()


@dataclass(frozen=True)
class Publication:
    index: int
    title: str
    context: str
    pdf_url: str
    cve: str


class AutomationError(RuntimeError):
    """Error controlado de la automatización."""


def session() -> requests.Session:
    current = requests.Session()
    current.headers.update({"User-Agent": USER_AGENT})
    return current


def clean_space(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip(" |-\n\t")


def fetch_text(http: requests.Session, url: str) -> str:
    response = http.get(url, timeout=REQUEST_TIMEOUT)
    response.raise_for_status()
    response.encoding = response.apparent_encoding or response.encoding
    return response.text


def fetch_bytes(http: requests.Session, url: str) -> bytes:
    response = http.get(url, timeout=REQUEST_TIMEOUT)
    response.raise_for_status()
    return response.content


def parse_spanish_date(text: str) -> date:
    match = re.search(
        r"(\d{1,2})\s+de\s+([A-Za-zÁÉÍÓÚÜÑáéíóúüñ]+)\s+de\s+(\d{4})",
        text,
        flags=re.IGNORECASE,
    )
    if not match:
        raise AutomationError(
            "No fue posible reconocer la fecha de la edición del Diario Oficial."
        )

    day = int(match.group(1))
    month_name = (
        match.group(2)
        .lower()
        .replace("á", "a")
        .replace("é", "e")
        .replace("í", "i")
        .replace("ó", "o")
        .replace("ú", "u")
        .replace("ü", "u")
    )
    month = MONTHS.get(month_name)
    if not month:
        raise AutomationError(f"Mes no reconocido: {match.group(2)}")

    return date(int(match.group(3)), month, day)


def extract_edition(html: str) -> tuple[str, date]:
    soup = BeautifulSoup(html, "html.parser")
    page_text = clean_space(soup.get_text(" ", strip=True))

    edition_match = re.search(
        r"Edici[oó]n\s+N[uú]m\.?\s*([0-9.]+)",
        page_text,
        flags=re.IGNORECASE,
    )
    if not edition_match:
        raise AutomationError("No fue posible reconocer el número de edición.")

    edition_number = edition_match.group(1)
    edition_date = parse_spanish_date(page_text)
    return edition_number, edition_date


def nearby_context(link: Any) -> str:
    pieces: list[str] = []

    parent = link.parent
    if parent:
        parent_text = clean_space(parent.get_text(" ", strip=True))
        if parent_text:
            pieces.append(parent_text)

    previous = link.find_previous(string=True)
    attempts = 0
    while previous and attempts < 6:
        text = clean_space(str(previous))
        if text and "Ver PDF" not in text:
            pieces.append(text)
        previous = previous.find_previous(string=True)
        attempts += 1

    combined = clean_space(" | ".join(dict.fromkeys(reversed(pieces))))
    combined = re.sub(
        r"\|\s*Ver PDF\s*\(CVE-[^)]+\)",
        "",
        combined,
        flags=re.IGNORECASE,
    )
    return combined[-1200:]


def extract_publications(html: str) -> list[Publication]:
    soup = BeautifulSoup(html, "html.parser")
    publications: list[Publication] = []
    seen_urls: set[str] = set()

    for link in soup.find_all("a", href=True):
        label = clean_space(link.get_text(" ", strip=True))
        href = str(link.get("href", "")).strip()

        is_pdf = ".pdf" in href.lower()
        is_publication = "ver pdf" in label.lower() or is_pdf
        if not is_publication:
            continue

        pdf_url = urljoin(INDEX_URL, href)
        if pdf_url in seen_urls:
            continue
        seen_urls.add(pdf_url)

        context = nearby_context(link)
        title = context

        # Reduce encabezados muy largos conservando la parte más cercana al enlace.
        if " | " in title:
            parts = [clean_space(part) for part in title.split(" | ") if clean_space(part)]
            if parts:
                title = parts[-1]

        title = re.sub(
            r"Ver PDF\s*\(CVE-[^)]+\)",
            "",
            title,
            flags=re.IGNORECASE,
        )
        title = clean_space(title) or f"Publicación CVE {label}"

        cve_match = re.search(r"CVE[-\s]?(\d+)", f"{label} {href}", flags=re.IGNORECASE)
        cve = cve_match.group(1) if cve_match else ""

        publications.append(
            Publication(
                index=len(publications),
                title=title,
                context=context,
                pdf_url=pdf_url,
                cve=cve,
            )
        )

    if not publications:
        raise AutomationError(
            "La edición fue encontrada, pero no se detectaron publicaciones PDF."
        )

    return publications


def extract_json_object(raw: str) -> dict[str, Any]:
    text = raw.strip()
    text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s*```$", "", text)

    try:
        parsed = json.loads(text)
    except json.JSONDecodeError:
        start = text.find("{")
        end = text.rfind("}")
        if start < 0 or end <= start:
            raise AutomationError("La respuesta de IA no contiene JSON válido.")
        try:
            parsed = json.loads(text[start : end + 1])
        except json.JSONDecodeError as exc:
            raise AutomationError(
                f"No fue posible interpretar la respuesta JSON de IA: {exc}"
            ) from exc

    if not isinstance(parsed, dict):
        raise AutomationError("La respuesta estructurada debe ser un objeto JSON.")
    return parsed


def openai_client() -> OpenAI:
    if not os.environ.get("OPENAI_API_KEY"):
        raise AutomationError(
            "Falta el secreto OPENAI_API_KEY en el repositorio de GitHub."
        )
    return OpenAI()


def classify_publications(
    client: OpenAI,
    publications: list[Publication],
) -> list[int]:
    compact = [
        {
            "index": item.index,
            "titulo": item.title,
            "contexto": item.context,
            "cve": item.cve,
        }
        for item in publications
    ]

    prompt = f"""
Eres un analista normativo del Departamento de Estudios Inmobiliarios de
Transsa. Revisa el sumario oficial del Diario Oficial de Chile y selecciona
las publicaciones que deben analizarse para el Observatorio Normativo Urbano.

CRITERIOS:
{RELEVANCE_CRITERIA}

PUBLICACIONES:
{json.dumps(compact, ensure_ascii=False, indent=2)}

Devuelve exclusivamente JSON válido con esta estructura:
{{
  "relevant_indices": [0, 2, 5]
}}

No inventes índices. Si no hay publicaciones pertinentes, devuelve una lista
vacía.
""".strip()

    response = client.responses.create(
        model=MODEL,
        input=prompt,
        store=False,
    )
    result = extract_json_object(response.output_text)
    indices = result.get("relevant_indices", [])

    if not isinstance(indices, list):
        raise AutomationError("relevant_indices no es una lista.")

    valid = sorted(
        {
            int(index)
            for index in indices
            if isinstance(index, (int, float, str))
            and str(index).strip().isdigit()
            and 0 <= int(index) < len(publications)
        }
    )
    return valid


def pdf_to_text(pdf_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(pdf_bytes))
    pages: list[str] = []

    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if text:
            pages.append(text)

    output = clean_space("\n".join(pages))
    return output[:MAX_PDF_CHARS]


def analyze_documents(
    client: OpenAI,
    documents: list[dict[str, Any]],
    edition_number: str,
    edition_date: date,
) -> list[dict[str, Any]]:
    source_material = [
        {
            "document_index": document["document_index"],
            "titulo_sumario": document["title"],
            "contexto_sumario": document["context"],
            "cve": document["cve"],
            "texto_pdf_oficial": document["text"],
        }
        for document in documents
    ]

    prompt = f"""
Actúa como analista normativo del Departamento de Estudios Inmobiliarios de
Transsa. Analiza únicamente los documentos oficiales entregados, publicados
en la edición N.º {edition_number} del Diario Oficial de Chile, de fecha
{edition_date.isoformat()}.

OBJETIVO:
Preparar registros breves, precisos y útiles para un portal interno de
seguimiento sobre planificación urbana, urbanismo, construcción, arquitectura,
vivienda, patrimonio y evaluación ambiental con incidencia territorial.

REGLAS:
- No inventes antecedentes.
- Distingue correctamente alcance Nacional, Regional o Comunal.
- La implicancia debe explicar qué cambia o qué etapa se abre en la práctica.
- Aclara cuando una publicación no modifica directamente un instrumento de
  planificación ni una norma urbanística.
- Si un documento finalmente no es pertinente, no lo incluyas.
- Usa nombres oficiales de organismos, regiones y comunas.
- Un mismo documento puede originar solo un registro, salvo que contenga dos
  modificaciones completamente independientes.

DOCUMENTOS:
{json.dumps(source_material, ensure_ascii=False, indent=2)}

Devuelve exclusivamente JSON válido con esta estructura:
{{
  "items": [
    {{
      "document_index": 0,
      "titulo": "Título ejecutivo y específico",
      "escala": "Nacional | Regional | Comunal",
      "categoria": "Planificación urbana",
      "region": "Nombre de la región o Chile",
      "comuna": "Nombre de la comuna o cadena vacía",
      "organismo": "Organismo emisor",
      "tipo_norma": "Ley, decreto, resolución, extracto de EIA, etc.",
      "numero": "Número o identificación del acto",
      "resumen": "Qué se publicó, en lenguaje claro.",
      "implicancia": "Consecuencia práctica y límites de la publicación.",
      "impactados": "Actores que podrían verse afectados."
    }}
  ]
}}
""".strip()

    response = client.responses.create(
        model=MODEL,
        input=prompt,
        store=False,
    )
    result = extract_json_object(response.output_text)
    items = result.get("items", [])

    if not isinstance(items, list):
        raise AutomationError("El campo items no es una lista.")

    documents_by_index = {
        int(document["document_index"]): document for document in documents
    }
    records: list[dict[str, Any]] = []

    for item in items:
        if not isinstance(item, dict):
            continue

        try:
            document_index = int(item.get("document_index"))
        except (TypeError, ValueError):
            continue

        source = documents_by_index.get(document_index)
        if not source:
            continue

        scale = clean_space(str(item.get("escala", "")))
        if scale not in {"Nacional", "Regional", "Comunal"}:
            scale = "Regional"

        record = {
            "fecha": edition_date.isoformat(),
            "titulo": clean_space(str(item.get("titulo", "")))
            or source["title"],
            "estado": "Con novedades",
            "escala": scale,
            "categoria": clean_space(str(item.get("categoria", "")))
            or "Normativa urbana",
            "region": clean_space(str(item.get("region", ""))) or "Chile",
            "comuna": clean_space(str(item.get("comuna", ""))),
            "organismo": clean_space(str(item.get("organismo", "")))
            or "Organismo no identificado",
            "tipo_norma": clean_space(str(item.get("tipo_norma", "")))
            or "Publicación oficial",
            "numero": clean_space(str(item.get("numero", ""))),
            "resumen": clean_space(str(item.get("resumen", ""))),
            "implicancia": clean_space(str(item.get("implicancia", ""))),
            "impactados": clean_space(str(item.get("impactados", ""))),
            "destacado": False,
            "source_url": source["pdf_url"],
            "cve": source["cve"],
            "edicion": edition_number,
        }

        if not record["resumen"]:
            continue
        records.append(record)

    return records


def load_reports() -> list[dict[str, Any]]:
    raw = DATA_FILE.read_text(encoding="utf-8").strip()
    prefix = "window.REPORTES = "

    if not raw.startswith(prefix) or not raw.endswith(";"):
        raise AutomationError(
            "data/reportes.js no tiene el formato esperado."
        )

    reports = json.loads(raw[len(prefix) : -1])
    if not isinstance(reports, list):
        raise AutomationError("La base de reportes no es una lista.")
    return reports


def save_reports(reports: list[dict[str, Any]]) -> None:
    reports.sort(
        key=lambda item: (
            str(item.get("fecha", "")),
            str(item.get("titulo", "")),
        ),
        reverse=True,
    )
    DATA_FILE.write_text(
        "window.REPORTES = "
        + json.dumps(reports, ensure_ascii=False, indent=2)
        + ";\n",
        encoding="utf-8",
    )


def has_date(reports: list[dict[str, Any]], target: date) -> bool:
    value = target.isoformat()
    return any(str(item.get("fecha", "")) == value for item in reports)


def no_news_record(
    report_date: date,
    edition_number: str,
    edition_date: date,
    reason: str,
) -> dict[str, Any]:
    return {
        "fecha": report_date.isoformat(),
        "titulo": "Sin cambios normativos relevantes",
        "estado": "Sin novedades",
        "escala": "Nacional",
        "categoria": "Sin novedades",
        "region": "Chile",
        "comuna": "",
        "organismo": "Diario Oficial de la República de Chile",
        "tipo_norma": "Revisión diaria",
        "numero": "",
        "resumen": reason,
        "implicancia": (
            "No corresponde actualizar matrices normativas ni bases de "
            "instrumentos de planificación territorial por esta revisión."
        ),
        "impactados": "Sin impacto normativo nuevo identificado.",
        "destacado": False,
        "source_url": INDEX_URL,
        "cve": "",
        "edicion": edition_number,
        "edicion_fecha": edition_date.isoformat(),
    }


def add_hyperlink_text(paragraph: Any, label: str, url: str) -> None:
    run = paragraph.add_run(f"{label}: {url}")
    run.font.name = "Outfit"
    run.font.size = Pt(9)


def generate_docx(report_date: date, records: list[dict[str, Any]]) -> Path:
    output_dir = (
        DOCUMENTS_DIR
        / f"{report_date.year:04d}"
        / f"{report_date.month:02d}"
    )
    output_dir.mkdir(parents=True, exist_ok=True)
    output = output_dir / (
        f"Reporte_normativo_urbano_{report_date.isoformat()}.docx"
    )

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = document.styles
    styles["Normal"].font.name = "Outfit"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Outfit"
    styles["Title"].font.size = Pt(23)
    styles["Title"].font.bold = True
    styles["Heading 1"].font.name = "Outfit"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.name = "Outfit"
    styles["Heading 2"].font.size = Pt(12)
    styles["Heading 2"].font.bold = True

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("Reporte normativo urbano")

    subtitle = document.add_paragraph()
    subtitle.add_run(
        f"Departamento de Estudios Inmobiliarios · Transsa\n"
        f"{report_date.strftime('%d-%m-%Y')}"
    ).bold = True

    document.add_heading("Resumen ejecutivo", level=1)
    with_changes = [r for r in records if r.get("estado") == "Con novedades"]

    if with_changes:
        document.add_paragraph(
            f"Se identificaron {len(with_changes)} publicaciones relevantes "
            "para planificación urbana, construcción, arquitectura o análisis "
            "territorial."
        )
    else:
        document.add_paragraph(
            records[0]["resumen"]
            if records
            else "No se identificaron novedades relevantes."
        )

    for scale in ("Nacional", "Regional", "Comunal"):
        scale_records = [r for r in records if r.get("escala") == scale]
        if not scale_records:
            continue

        document.add_heading(f"Alcance {scale.lower()}", level=1)

        for record in scale_records:
            document.add_heading(record["titulo"], level=2)
            details = [
                ("Publicación", record.get("tipo_norma", "")),
                ("Número", record.get("numero", "")),
                ("Organismo", record.get("organismo", "")),
                ("Región", record.get("region", "")),
                ("Comuna", record.get("comuna", "")),
            ]

            for label, value in details:
                if value:
                    paragraph = document.add_paragraph()
                    paragraph.add_run(f"{label}: ").bold = True
                    paragraph.add_run(str(value))

            paragraph = document.add_paragraph()
            paragraph.add_run("Qué se publicó: ").bold = True
            paragraph.add_run(record.get("resumen", ""))

            paragraph = document.add_paragraph()
            paragraph.add_run("Implicancia práctica: ").bold = True
            paragraph.add_run(record.get("implicancia", ""))

            paragraph = document.add_paragraph()
            paragraph.add_run("Actores impactados: ").bold = True
            paragraph.add_run(record.get("impactados", ""))

            source_url = record.get("source_url", "")
            if source_url:
                add_hyperlink_text(
                    document.add_paragraph(),
                    "Fuente oficial",
                    source_url,
                )

    document.add_paragraph()
    footer = document.add_paragraph(
        "Documento generado automáticamente para uso interno. "
        "La fuente oficial y jurídicamente válida es el Diario Oficial "
        "de la República de Chile."
    )
    footer.runs[0].italic = True
    footer.runs[0].font.size = Pt(8.5)

    document.save(output)
    return output


def attach_word_url(
    records: list[dict[str, Any]],
    docx_path: Path,
) -> None:
    relative = docx_path.relative_to(ROOT).as_posix()
    for record in records:
        record["word_url"] = relative


def append_unique(
    existing: list[dict[str, Any]],
    incoming: list[dict[str, Any]],
) -> int:
    keys = {
        (
            str(item.get("fecha", "")),
            clean_space(str(item.get("titulo", ""))).lower(),
        )
        for item in existing
    }

    added = 0
    for item in incoming:
        key = (
            str(item.get("fecha", "")),
            clean_space(str(item.get("titulo", ""))).lower(),
        )
        if key in keys:
            continue
        existing.append(item)
        keys.add(key)
        added += 1

    return added


def main() -> int:
    today = datetime.now(TIMEZONE).date()
    http = session()
    reports = load_reports()

    print(f"Fecha de ejecución en Chile: {today.isoformat()}")
    print("Descargando sumario oficial...")
    html = fetch_text(http, INDEX_URL)

    edition_number, edition_date = extract_edition(html)
    print(
        f"Última edición detectada: N.º {edition_number}, "
        f"{edition_date.isoformat()}"
    )

    # En domingos, festivos o antes de una nueva publicación, el sitio puede
    # mantener como última edición una fecha anterior.
    if edition_date < today:
        if has_date(reports, today):
            print("La revisión de hoy ya existe. No se realizan cambios.")
            return 0

        record = no_news_record(
            report_date=today,
            edition_number=edition_number,
            edition_date=edition_date,
            reason=(
                "Al momento de la revisión no se encontraba publicada una "
                "nueva edición del Diario Oficial para esta fecha. La edición "
                f"más reciente disponible correspondía al {edition_date.isoformat()}."
            ),
        )
        records = [record]
        docx = generate_docx(today, records)
        attach_word_url(records, docx)
        append_unique(reports, records)
        save_reports(reports)
        print(f"Registro sin nueva edición agregado: {today.isoformat()}")
        return 0

    if edition_date > today:
        raise AutomationError(
            "La fecha detectada en el Diario Oficial es posterior a la fecha "
            "actual en Chile. Se detiene para evitar registros incorrectos."
        )

    if has_date(reports, edition_date):
        print("La edición de hoy ya fue procesada. No se realizan cambios.")
        return 0

    publications = extract_publications(html)
    print(f"Publicaciones PDF detectadas: {len(publications)}")

    client = openai_client()
    relevant_indices = classify_publications(client, publications)
    print(f"Publicaciones preseleccionadas: {len(relevant_indices)}")

    records: list[dict[str, Any]]

    if relevant_indices:
        documents: list[dict[str, Any]] = []

        for index in relevant_indices:
            publication = publications[index]
            print(f"Descargando CVE {publication.cve or 'sin CVE'}...")
            pdf_bytes = fetch_bytes(http, publication.pdf_url)
            pdf_text = pdf_to_text(pdf_bytes)

            if not pdf_text:
                print(
                    f"Advertencia: no se pudo extraer texto del PDF "
                    f"{publication.pdf_url}",
                    file=sys.stderr,
                )
                continue

            documents.append(
                {
                    "document_index": publication.index,
                    "title": publication.title,
                    "context": publication.context,
                    "pdf_url": publication.pdf_url,
                    "cve": publication.cve,
                    "text": pdf_text,
                }
            )

        records = (
            analyze_documents(
                client=client,
                documents=documents,
                edition_number=edition_number,
                edition_date=edition_date,
            )
            if documents
            else []
        )
    else:
        records = []

    if not records:
        records = [
            no_news_record(
                report_date=edition_date,
                edition_number=edition_number,
                edition_date=edition_date,
                reason=(
                    "Se revisó la edición más reciente del Diario Oficial y "
                    "no se identificaron cambios normativos relevantes para "
                    "planificación urbana, construcción, arquitectura, "
                    "vivienda, patrimonio o evaluación ambiental territorial."
                ),
            )
        ]

    docx = generate_docx(edition_date, records)
    attach_word_url(records, docx)

    added = append_unique(reports, records)
    if added:
        save_reports(reports)

    print(f"Registros agregados: {added}")
    print(f"Word diario: {docx.relative_to(ROOT).as_posix()}")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except requests.RequestException as exc:
        print(f"Error de conexión: {exc}", file=sys.stderr)
        raise SystemExit(2)
    except AutomationError as exc:
        print(f"Error controlado: {exc}", file=sys.stderr)
        raise SystemExit(3)
    except Exception as exc:
        print(f"Error inesperado: {type(exc).__name__}: {exc}", file=sys.stderr)
        raise
