from __future__ import annotations

from datetime import date, datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from core.models import CanonicalEvent
from core.vocabulary import SPANISH_LABELS

MONTHS_ES = [
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]

def clean_edition_number(value: str) -> str:
    return str(value or "").strip().rstrip(". ·")


def format_date_es(value: str | date) -> str:
    if isinstance(value, date):
        parsed = value
    else:
        try:
            parsed = datetime.fromisoformat(str(value)).date()
        except (TypeError, ValueError):
            return str(value or "")
    return f"{parsed.day} de {MONTHS_ES[parsed.month - 1]} de {parsed.year}"

def official_fields(event: CanonicalEvent) -> dict:
    payload = event.legacy_payload if isinstance(event.legacy_payload, dict) else {}
    fields = payload.get("official_act") or {}
    return fields if isinstance(fields, dict) else {}

def territorial_scope_label(event: CanonicalEvent) -> str:
    if event.territory.region and event.territory.commune:
        return "Regional y comunal"
    return SPANISH_LABELS.get(event.territory.scale, event.territory.scale or "Sin determinar")


def territory_label(event: CanonicalEvent) -> tuple[str, str, str]:
    territory = event.territory
    scale = SPANISH_LABELS.get(territory.scale, territory.scale or "Sin determinar")
    region = territory.region or ("Chile" if territory.scale == "national" else "")
    commune = territory.commune
    return scale, region, commune


def event_to_legacy_report(
    event: CanonicalEvent,
    word_path: Path | None = None,
    root: Path | None = None,
) -> dict:
    scale, region, commune = territory_label(event)
    source = event.source
    result = {
        "fecha": event.event_date,
        "titulo": event.title,
        "estado": "Con novedades",
        "escala": scale,
        "categoria": event.category.replace("_", " ").title(),
        "region": region or "Chile",
        "comuna": commune,
        "organismo": source.source_name,
        "tipo_norma": source.document_type or SPANISH_LABELS.get(event.event_type, event.event_type),
        "numero": source.document_number,
        "resumen": event.summary,
        "implicancia": event.practical_implications or event.why_it_matters,
        "impactados": event.impacted_parties,
        "destacado": event.is_featured,
        "source_url": source.url,
        "cve": source.external_id,
        "edicion": source.edition,
        "estado_revision": event.review_status,
        "accion_recomendada": event.recommended_action,
        "event_id": event.event_id,
    }
    if word_path is not None:
        result["word_url"] = (
            word_path.relative_to(root).as_posix()
            if root is not None
            else word_path.as_posix()
        )
    return result


def no_news_legacy_report(
    report_date: date,
    edition_number: str,
    edition_date: date,
    reason: str,
    word_path: Path | None = None,
    root: Path | None = None,
) -> dict:
    result = {
        "fecha": report_date.isoformat(),
        "titulo": "Sin cambios relevantes identificados",
        "estado": "Sin novedades",
        "escala": "Nacional",
        "categoria": "Sin novedades",
        "region": "Chile",
        "comuna": "",
        "organismo": "Diario Oficial de la República de Chile",
        "tipo_norma": "Revisión diaria",
        "numero": "",
        "resumen": reason,
        "implicancia": (
            "No se genera una acción automática. El resultado queda registrado "
            "para mantener la trazabilidad diaria."
        ),
        "impactados": "Sin impacto nuevo identificado.",
        "destacado": False,
        "source_url": "https://www.diariooficial.interior.gob.cl/edicionelectronica/index.php/bom.php",
        "cve": "",
        "edicion": edition_number,
        "edicion_fecha": edition_date.isoformat(),
        "estado_revision": "preliminary",
        "accion_recomendada": "Sin acción requerida.",
    }
    if word_path is not None:
        result["word_url"] = (
            word_path.relative_to(root).as_posix()
            if root is not None
            else word_path.as_posix()
        )
    return result


def generate_daily_docx(
    report_date: date,
    edition_number: str,
    events: Iterable[CanonicalEvent],
    output_path: Path,
    no_news_reason: str = "",
) -> Path:
    events = list(events)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(22)
    styles["Title"].font.bold = True
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.name = "Aptos Display"
    styles["Heading 2"].font.size = Pt(12)
    styles["Heading 2"].font.bold = True

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("Transsa Urban Intelligence")

    subtitle = document.add_paragraph()
    run = subtitle.add_run(
        "Reporte preliminar del Diario Oficial\n"
        f"Edición N.º {clean_edition_number(edition_number)} · {format_date_es(report_date)}"
    )
    run.bold = True

    document.add_heading("Resumen ejecutivo", level=1)
    if events:
        count = len(events)
        noun = "evento preliminar relevante" if count == 1 else "eventos preliminares relevantes"
        document.add_paragraph(
            f"Se identificaron {count} {noun} para normativa territorial, desarrollo urbano, "
            "construcción, vivienda, infraestructura o mercado inmobiliario. "
            "Los antecedentes deben ser revisados antes de marcarse como validados."
        )
    else:
        document.add_paragraph(no_news_reason or "No se identificaron novedades relevantes.")

    grouped: dict[str, list[CanonicalEvent]] = {}
    for event in events:
        key = "regional_communal" if event.territory.region and event.territory.commune else (event.territory.scale or "undetermined")
        grouped.setdefault(key, []).append(event)

    scale_order = [
        "national",
        "interregional",
        "regional_communal",
        "regional",
        "provincial",
        "intercommunal",
        "communal",
        "local",
        "multiple",
        "undetermined",
    ]

    for scale in scale_order:
        scale_events = grouped.get(scale, [])
        if not scale_events:
            continue
        heading_label = "regional y comunal" if scale == "regional_communal" else SPANISH_LABELS.get(scale, scale).lower()
        document.add_heading(f"Alcance {heading_label}", level=1)
        for event in scale_events:
            document.add_heading(event.title, level=2)
            official = official_fields(event)
            publication_source = event.legacy_payload.get("publication_source", "Diario Oficial de la República de Chile") if isinstance(event.legacy_payload, dict) else "Diario Oficial de la República de Chile"
            details = [
                ("Tipo", SPANISH_LABELS.get(event.event_type, event.event_type)),
                ("Escala", territorial_scope_label(event)),
                ("Organismo emisor", official.get("official_issuer") or event.source.source_name),
                ("Fuente de publicación", publication_source),
                ("Tipo de acto", official.get("act_type") or event.source.document_type),
                ("Número del acto", official.get("act_number") or event.source.document_number),
                ("Fecha del acto", format_date_es(official.get("act_date")) if official.get("act_date") else ""),
                ("Etapa del procedimiento", official.get("procedure_stage", "")),
                ("Plazo de participación", f"{official.get('participation_days')} días hábiles" if official.get("participation_days") else ""),
                ("Inicio del cómputo del plazo", official.get("participation_start_rule", "")),
                ("Base legal", official.get("legal_basis", "")),
                ("Proyecto", official.get("project_name", "")),
                ("Titular o proponente", official.get("project_holder", "")),
                ("Antecedentes concretos del proyecto", official.get("project_description", "")),
                ("Región", event.territory.region),
                ("Provincia", event.territory.province),
                ("Comuna", event.territory.commune),
                ("Ubicación específica", official.get("location_detail", "")),
                ("Relevancia", {"low": "Baja", "medium": "Media", "high": "Alta", "critical": "Crítica"}.get(event.relevance_level, event.relevance_level)),
                ("Estado", SPANISH_LABELS.get(event.review_status, event.review_status)),
            ]
            for label, value in details:
                if value:
                    paragraph = document.add_paragraph()
                    paragraph.add_run(f"{label}: ").bold = True
                    paragraph.add_run(str(value))

            fields = [
                ("Qué ocurrió", event.summary),
                ("Por qué importa", event.why_it_matters),
                ("Implicancias prácticas", event.practical_implications),
                ("Actores impactados", event.impacted_parties),
                ("Acción sugerida", event.recommended_action),
            ]
            for label, value in fields:
                if value:
                    paragraph = document.add_paragraph()
                    paragraph.add_run(f"{label}: ").bold = True
                    paragraph.add_run(value)

            if event.requires_review_reason:
                paragraph = document.add_paragraph()
                paragraph.add_run("Revisión requerida: ").bold = True
                paragraph.add_run(event.requires_review_reason)

            if event.source.url:
                paragraph = document.add_paragraph()
                paragraph.add_run("Fuente oficial: ").bold = True
                paragraph.add_run(event.source.url)

    footer = document.add_paragraph()
    run = footer.add_run(
        "Reporte generado localmente con reglas y Ollama. Su estado es preliminar; "
        "la fuente oficial mantiene plena validez y la revisión humana es obligatoria."
    )
    run.italic = True
    run.font.size = Pt(8.5)

    document.save(output_path)
    return output_path
