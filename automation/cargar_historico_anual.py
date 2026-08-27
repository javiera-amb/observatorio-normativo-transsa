#!/usr/bin/env python3
"""
Carga un histórico anual de normativa urbana e IPT usando fuentes oficiales.

Por defecto procesa el año 2025. Puede ejecutarse para todo el año o para un
mes específico mediante variables de entorno:

- HISTORIC_YEAR=2025
- HISTORIC_MONTH=all | 1..12
- HISTORIC_FORCE=true | false

Resultados:
- data/historicos.js
- documentos/historicos/Reporte_anual_normativo_YYYY.docx
- consolidados/historicos/Consolidado_anual_normativo_YYYY.csv

La búsqueda automática depende de la indexación de fuentes oficiales y debe
ser validada antes de usar los resultados en decisiones normativas.
"""

from __future__ import annotations

import csv
import json
import os
import re
import sys
from datetime import date, datetime
from pathlib import Path
from typing import Any
from zoneinfo import ZoneInfo

from docx import Document
from docx.shared import Cm, Pt
from openai import OpenAI


ROOT = Path(__file__).resolve().parents[1]
DATA_FILE = ROOT / "data" / "historicos.js"
DOCS_DIR = ROOT / "documentos" / "historicos"
CSV_DIR = ROOT / "consolidados" / "historicos"
TIMEZONE = ZoneInfo("America/Santiago")
MODEL = os.environ.get("OPENAI_MODEL", "gpt-5-mini")

OFFICIAL_DOMAINS = [
    "diariooficial.interior.gob.cl",
    "minvu.gob.cl",
    "ide.minvu.cl",
    "eae.mma.gob.cl",
    "mma.gob.cl",
    "sea.gob.cl",
    "gob.cl",
    "bcn.cl",
]

MONTH_NAMES = [
    "", "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]

REQUIRED_FIELDS = [
    "periodo", "modulo", "fecha", "region", "comuna", "escala",
    "categoria", "tipo_norma", "numero", "organismo", "titulo",
    "resumen", "implicancia", "estado", "fuente",
]


class HistoricalError(RuntimeError):
    pass


def clean(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def parse_json(raw: str) -> dict[str, Any]:
    text = raw.strip()
    text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s*```$", "", text)

    try:
        value = json.loads(text)
    except json.JSONDecodeError:
        start = text.find("{")
        end = text.rfind("}")
        if start < 0 or end <= start:
            raise HistoricalError("La respuesta no contiene JSON válido.")
        value = json.loads(text[start:end + 1])

    if not isinstance(value, dict):
        raise HistoricalError("La respuesta debe ser un objeto JSON.")
    return value


def api_client() -> OpenAI:
    if not os.environ.get("OPENAI_API_KEY"):
        raise HistoricalError("Falta OPENAI_API_KEY en GitHub.")
    return OpenAI()


def month_range(year: int, month: int) -> tuple[date, date]:
    start = date(year, month, 1)
    if month == 12:
        next_month = date(year + 1, 1, 1)
    else:
        next_month = date(year, month + 1, 1)
    end = date.fromordinal(next_month.toordinal() - 1)
    return start, end


def research_month(
    client: OpenAI,
    year: int,
    month: int,
) -> list[dict[str, Any]]:
    start, end = month_range(year, month)
    period = f"{year:04d}-{month:02d}"

    prompt = f"""
Eres analista normativo del Departamento de Estudios Inmobiliarios de Transsa.
Construye el histórico oficial de Chile para el período {start.isoformat()} a
{end.isoformat()}.

Debes buscar DOS grupos:

A. DIARIO OFICIAL
Publicaciones relevantes sobre:
- planificación urbana y territorial;
- LGUC, OGUC y normas de construcción;
- permisos, recepciones, subdivisiones y urbanizaciones;
- vivienda y subsidios cuando alteren reglas, llamados o montos;
- patrimonio arquitectónico y arqueológico vinculado a obras;
- evaluación ambiental de proyectos con incidencia territorial;
- infraestructura, movilidad, expropiaciones y espacio público.

B. INSTRUMENTOS DE PLANIFICACIÓN TERRITORIAL
Actos vigentes o en tramitación sobre:
- PRC, PRI, PRM, límites urbanos;
- planes seccionales y enmiendas;
- postergaciones de permisos;
- PROT y evaluación ambiental estratégica de IPT;
- aprobaciones, promulgaciones, modificaciones y rectificaciones.

Revisa Chile completo y prioriza fuentes oficiales. No inventes actos ni URLs.
No incluyas noticias genéricas sin acto, publicación o etapa identificable.
Ante duda razonable, incluye el registro y explica su estado.

Devuelve exclusivamente JSON válido:
{{
  "items": [
    {{
      "periodo": "{period}",
      "modulo": "Diario Oficial | IPT",
      "fecha": "YYYY-MM-DD o vacío",
      "region": "Región o Chile",
      "comuna": "Comuna o vacío",
      "escala": "Nacional | Regional | Comunal | Intercomunal",
      "categoria": "Materia principal",
      "tipo_norma": "Ley, decreto, resolución, PRC, enmienda, etc.",
      "numero": "Número o identificación",
      "organismo": "Organismo oficial",
      "titulo": "Título ejecutivo",
      "resumen": "Qué ocurrió",
      "implicancia": "Consecuencia práctica y límites",
      "estado": "Vigente | En tramitación | Consulta | Publicado | Otro",
      "fuente": "URL oficial completa"
    }}
  ],
  "nota_cobertura": "Advertencias sobre fuentes o exhaustividad"
}}
""".strip()

    response = client.responses.create(
        model=MODEL,
        input=prompt,
        tools=[
            {
                "type": "web_search",
                "filters": {"allowed_domains": OFFICIAL_DOMAINS},
                "search_context_size": "high",
                "user_location": {
                    "type": "approximate",
                    "country": "CL",
                    "timezone": "America/Santiago",
                },
            }
        ],
        store=False,
    )

    result = parse_json(response.output_text)
    items = result.get("items", [])
    if not isinstance(items, list):
        return []

    output: list[dict[str, Any]] = []
    for item in items:
        if not isinstance(item, dict):
            continue

        record = {field: clean(item.get(field, "")) for field in REQUIRED_FIELDS}
        record["periodo"] = period

        if record["modulo"] not in {"Diario Oficial", "IPT"}:
            record["modulo"] = "Diario Oficial"

        if not record["titulo"] or not record["resumen"] or not record["fuente"]:
            continue

        output.append(record)

    return output


def deduplicate(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    seen: set[tuple[str, ...]] = set()
    result: list[dict[str, Any]] = []

    for item in items:
        key = (
            item.get("periodo", "").lower(),
            item.get("modulo", "").lower(),
            item.get("fecha", "").lower(),
            item.get("region", "").lower(),
            item.get("comuna", "").lower(),
            item.get("titulo", "").lower(),
            item.get("numero", "").lower(),
        )
        if key in seen:
            continue
        seen.add(key)
        result.append(item)

    return sorted(
        result,
        key=lambda item: (
            item.get("periodo", ""),
            item.get("fecha", ""),
            item.get("region", ""),
            item.get("comuna", ""),
        ),
        reverse=True,
    )


def load_reports() -> list[dict[str, Any]]:
    raw = DATA_FILE.read_text(encoding="utf-8").strip()
    prefix = "window.HISTORICOS = "
    if not raw.startswith(prefix) or not raw.endswith(";"):
        raise HistoricalError("Formato inválido en data/historicos.js.")
    value = json.loads(raw[len(prefix):-1])
    if not isinstance(value, list):
        raise HistoricalError("La base histórica no es una lista.")
    return value


def save_reports(reports: list[dict[str, Any]]) -> None:
    reports.sort(key=lambda report: int(report.get("year", 0)), reverse=True)
    DATA_FILE.write_text(
        "window.HISTORICOS = "
        + json.dumps(reports, ensure_ascii=False, indent=2)
        + ";\n",
        encoding="utf-8",
    )


def create_csv(year: int, items: list[dict[str, Any]]) -> Path:
    CSV_DIR.mkdir(parents=True, exist_ok=True)
    output = CSV_DIR / f"Consolidado_anual_normativo_{year}.csv"

    headers = [
        "Período", "Módulo", "Fecha", "Región", "Comuna", "Escala",
        "Categoría", "Tipo de norma o IPT", "Número", "Organismo",
        "Título", "Resumen", "Implicancia", "Estado", "Fuente oficial",
    ]

    with output.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle, delimiter=";")
        writer.writerow(headers)
        for item in items:
            writer.writerow([
                item.get("periodo", ""),
                item.get("modulo", ""),
                item.get("fecha", ""),
                item.get("region", ""),
                item.get("comuna", ""),
                item.get("escala", ""),
                item.get("categoria", ""),
                item.get("tipo_norma", ""),
                item.get("numero", ""),
                item.get("organismo", ""),
                item.get("titulo", ""),
                item.get("resumen", ""),
                item.get("implicancia", ""),
                item.get("estado", ""),
                item.get("fuente", ""),
            ])

    return output


def create_word(year: int, items: list[dict[str, Any]]) -> Path:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    output = DOCS_DIR / f"Reporte_anual_normativo_{year}.docx"

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)

    document.styles["Normal"].font.name = "Outfit"
    document.styles["Normal"].font.size = Pt(10)
    document.styles["Title"].font.name = "Outfit"
    document.styles["Title"].font.size = Pt(22)
    document.styles["Heading 1"].font.name = "Outfit"
    document.styles["Heading 1"].font.size = Pt(15)
    document.styles["Heading 2"].font.name = "Outfit"
    document.styles["Heading 2"].font.size = Pt(12)

    document.add_heading(f"Reporte normativo anual {year}", 0)
    document.add_paragraph(
        "Departamento de Estudios Inmobiliarios · Transsa"
    ).runs[0].bold = True

    communes = len({item.get("comuna") for item in items if item.get("comuna")})
    regions = len({item.get("region") for item in items if item.get("region")})
    daily_count = sum(1 for item in items if item.get("modulo") == "Diario Oficial")
    ipt_count = sum(1 for item in items if item.get("modulo") == "IPT")

    document.add_heading("Síntesis ejecutiva", level=1)
    document.add_paragraph(
        f"El consolidado contiene {len(items)} registros: "
        f"{daily_count} del Diario Oficial y {ipt_count} relacionados con IPT. "
        f"Se identificaron {communes} comunas y {regions} regiones."
    )

    for month in range(1, 13):
        period = f"{year:04d}-{month:02d}"
        month_items = [item for item in items if item.get("periodo") == period]
        if not month_items:
            continue

        document.add_heading(
            f"{MONTH_NAMES[month].capitalize()} {year}",
            level=1,
        )

        for item in month_items:
            document.add_heading(item.get("titulo", "Actualización"), level=2)

            for label, key in [
                ("Módulo", "modulo"),
                ("Fecha", "fecha"),
                ("Región", "region"),
                ("Comuna", "comuna"),
                ("Estado", "estado"),
                ("Organismo", "organismo"),
                ("Tipo", "tipo_norma"),
                ("Número", "numero"),
            ]:
                value = item.get(key, "")
                if value:
                    paragraph = document.add_paragraph()
                    paragraph.add_run(f"{label}: ").bold = True
                    paragraph.add_run(value)

            paragraph = document.add_paragraph()
            paragraph.add_run("Resumen: ").bold = True
            paragraph.add_run(item.get("resumen", ""))

            paragraph = document.add_paragraph()
            paragraph.add_run("Implicancia: ").bold = True
            paragraph.add_run(item.get("implicancia", ""))

            paragraph = document.add_paragraph()
            paragraph.add_run("Fuente oficial: ").bold = True
            paragraph.add_run(item.get("fuente", ""))

    note = document.add_paragraph(
        "La carga histórica automática depende de la indexación pública de "
        "las fuentes oficiales. Verifica cada antecedente en su fuente antes "
        "de aplicarlo a un análisis normativo o inmobiliario."
    )
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(8.5)

    document.save(output)
    return output


def build_summary(year: int, items: list[dict[str, Any]]) -> str:
    communes = len({item.get("comuna") for item in items if item.get("comuna")})
    regions = len({item.get("region") for item in items if item.get("region")})
    return (
        f"El histórico {year} contiene {len(items)} registros oficiales, "
        f"con referencias a {communes} comunas y {regions} regiones."
    )


def main() -> int:
    year = int(os.environ.get("HISTORIC_YEAR", "2025"))
    month_arg = clean(os.environ.get("HISTORIC_MONTH", "all")).lower()
    force = clean(os.environ.get("HISTORIC_FORCE", "false")).lower() == "true"

    if year < 2000 or year > datetime.now(TIMEZONE).year:
        raise HistoricalError("Año fuera de rango.")

    if month_arg == "all":
        months = list(range(1, 13))
    else:
        month = int(month_arg)
        if not 1 <= month <= 12:
            raise HistoricalError("El mes debe ser all o un número entre 1 y 12.")
        months = [month]

    reports = load_reports()
    report = next((item for item in reports if int(item.get("year", 0)) == year), None)

    if report is None:
        report = {
            "year": year,
            "titulo": f"Reporte anual normativo {year}",
            "fecha_generacion": "",
            "resumen_ejecutivo": "",
            "items": [],
            "meses_procesados": [],
            "word_url": "",
            "csv_url": "",
            "nota_cobertura": (
                "La revisión depende de la indexación pública de las fuentes "
                "oficiales y requiere validación profesional."
            ),
        }
        reports.append(report)

    existing_items = report.get("items", [])
    processed = set(int(value) for value in report.get("meses_procesados", []))

    api = api_client()

    for month in months:
        if month in processed and not force:
            print(f"{year}-{month:02d} ya estaba procesado. Se omite.")
            continue

        print(f"Investigando {MONTH_NAMES[month]} de {year}...")
        month_items = research_month(api, year, month)

        if force:
            period = f"{year:04d}-{month:02d}"
            existing_items = [
                item for item in existing_items
                if item.get("periodo") != period
            ]

        existing_items.extend(month_items)
        processed.add(month)

    items = deduplicate(existing_items)
    word = create_word(year, items)
    csv_file = create_csv(year, items)

    report.update({
        "fecha_generacion": datetime.now(TIMEZONE).strftime("%Y-%m-%d %H:%M"),
        "resumen_ejecutivo": build_summary(year, items),
        "items": items,
        "meses_procesados": sorted(processed),
        "word_url": word.relative_to(ROOT).as_posix(),
        "csv_url": csv_file.relative_to(ROOT).as_posix(),
    })

    save_reports(reports)

    print(f"Año: {year}")
    print(f"Meses procesados: {sorted(processed)}")
    print(f"Registros: {len(items)}")
    print(f"Word: {report['word_url']}")
    print(f"CSV: {report['csv_url']}")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except HistoricalError as exc:
        print(f"Error: {exc}", file=sys.stderr)
        raise SystemExit(2)
